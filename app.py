import streamlit as st
import tempfile
import os
import asyncio
import uuid
import re
import warnings
import nest_asyncio
import logging

# Configure logging according to ADK documentation
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(name)s - %(message)s'
)

logger = logging.getLogger(__name__)

# Apply nest_asyncio to allow nested event loops
nest_asyncio.apply()
logger.info("Applied nest_asyncio for nested event loop support")

# Suppress aiohttp warnings about event loop closure
warnings.filterwarnings('ignore', category=RuntimeWarning, message='.*Event loop is closed.*')
warnings.filterwarnings('ignore', category=RuntimeWarning, message='.*coroutine.*was never awaited.*')

# Apply compatibility shims before importing ADK modules
try:
    from my_agent.compat import aiohttp_compat
except ImportError:
    pass

from google.adk.runners import InMemoryRunner
from google.genai import types
from google.genai import Client
from docx import Document
from io import BytesIO
from my_agent.agent import sequential_agent, run_qna_query_simple as run_qna_query

logger.info("Imported agents and dependencies successfully")

# Helper to run async code
def run_sync(coro):
    """Run async function synchronously with nest_asyncio support"""
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

    return loop.run_until_complete(coro)


# Initialize runners and sessions once
@st.cache_resource
def initialize_runners():
    """Initialize ADK runners and create sessions"""
    logger.info("Initializing ADK runners and sessions")

    pipeline_runner = InMemoryRunner(
        agent=sequential_agent,
        app_name="segmentation_pipeline_improved"
    )

    # Create sessions
    user_id = "user1"
    pipeline_session_id = str(uuid.uuid4())

    logger.info(f"Created session ID: {pipeline_session_id}")

    # Create sessions using the session service
    run_sync(pipeline_runner.session_service.create_session(
        app_name="segmentation_pipeline_improved",
        user_id=user_id,
        session_id=pipeline_session_id
    ))

    logger.info("ADK runners and sessions initialized successfully")

    return {
        "pipeline_runner": pipeline_runner,
        "pipeline_session_id": pipeline_session_id,
        "user_id": user_id
    }


# Get initialized runners
runners = initialize_runners()

# Initialize session state
if 'clustering_results' not in st.session_state:
    st.session_state.clustering_results = None
if 'marketing_strategies' not in st.session_state:
    st.session_state.marketing_strategies = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'clustered_csv_path' not in st.session_state:
    st.session_state.clustered_csv_path = None
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'qna_context_initialized' not in st.session_state:
    st.session_state.qna_context_initialized = False

# Streamlit UI
st.title("Market Segmentation Agent")

# Step 1: File upload
uploaded_file = st.file_uploader("Upload a CSV file to perform market segmentation analysis", type="csv")

# Step 2: Run Analysis button
if uploaded_file and st.button("Run Analysis", type="primary"):
    logger.info("User initiated analysis")

    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    logger.info(f"Saved uploaded file to temporary path: {tmp_path}")

    with st.spinner("Running segmentation pipeline... This may take a few minutes."):
        try:
            logger.info("Starting segmentation pipeline")

            # Create Content object for the message
            content = types.Content(
                role='user',
                parts=[types.Part(text=f'Perform a complete market segmentation analysis and strategy for this data: "{tmp_path}"')]
            )

            # Run the pipeline and collect all events
            result_text = []
            events = runners["pipeline_runner"].run(
                user_id=runners["user_id"],
                session_id=runners["pipeline_session_id"],
                new_message=content
            )

            logger.debug("Processing pipeline events")

            # Process events
            clustering_text = []
            marketing_text = []
            current_agent = None
            event_count = 0

            for event in events:
                event_count += 1

                # Track which agent is responding
                if hasattr(event, 'agent_name') and event.agent_name:
                    current_agent = event.agent_name
                elif hasattr(event, 'name') and event.name:
                    current_agent = event.name

                # Check for agent info in event attributes
                event_str = str(event)
                if 'marketing' in event_str.lower():
                    current_agent = 'marketing_strategy_agent'
                elif 'clustering' in event_str.lower() and 'marketing' not in event_str.lower():
                    current_agent = 'clustering_agent'

                if hasattr(event, 'content') and event.content and event.content.parts:
                    for part in event.content.parts:
                        if hasattr(part, 'text') and part.text:
                            text = part.text
                            result_text.append(text)

                            # Split results by agent
                            if current_agent and 'clustering' in current_agent.lower():
                                clustering_text.append(text)
                            elif current_agent and 'marketing' in current_agent.lower():
                                marketing_text.append(text)

                # Extract state
                if hasattr(event, 'state') and event.state:
                    if isinstance(event.state, dict):
                        if 'clustering_results' in event.state:
                            st.session_state.clustering_results = event.state['clustering_results']
                        if 'marketing_strategies' in event.state:
                            st.session_state.marketing_strategies = event.state['marketing_strategies']

            logger.info(f"Processed {event_count} events from pipeline")

            result = '\n'.join(result_text)

            logger.info(f"Captured {len(clustering_text)} clustering text parts, {len(marketing_text)} marketing text parts")

            # Fallback: use text result if state not captured
            if not st.session_state.clustering_results:
                logger.warning("Clustering results not found in state, using fallback extraction")

                if clustering_text:
                    st.session_state.clustering_results = '\n'.join(clustering_text)
                elif result:
                    # Try to split by keywords
                    split_markers = ['Marketing Strateg', '## Marketing', '# Marketing', 'marketing strateg']
                    for marker in split_markers:
                        parts = result.split(marker, 1)
                        if len(parts) == 2:
                            st.session_state.clustering_results = parts[0]
                            st.session_state.marketing_strategies = marker + parts[1]
                            logger.info(f"Split results using marker: {marker}")
                            break
                    else:
                        st.session_state.clustering_results = result

            if not st.session_state.marketing_strategies and marketing_text:
                st.session_state.marketing_strategies = '\n'.join(marketing_text)

            # Final fallback for marketing strategies
            if not st.session_state.marketing_strategies and result:
                for marker in ['Marketing Strateg', '## Marketing', '# Marketing']:
                    if marker in result:
                        idx = result.index(marker)
                        st.session_state.marketing_strategies = result[idx:]
                        if st.session_state.clustering_results == result:
                            st.session_state.clustering_results = result[:idx]
                        break

            # Extract the path to the clustered CSV file from results
            try:
                if st.session_state.clustering_results:
                    results_str = str(st.session_state.clustering_results)
                    # Look for output_path in various formats
                    patterns = [
                        r"'output_path':\s*'([^']+)'",
                        r'"output_path":\s*"([^"]+)"',
                        r"output_path[:\s]+([^\s,\n]+\.csv)",
                        r"saved.*?to[:\s]+([^\s,\n]+\.csv)",
                        r"Path[:\s]+([^\s,\n]+\.csv)",
                        r"(/[\w/\-_.]+\.csv)",
                        r"([A-Z]:[\\\/][\w\\/\-_.]+\.csv)",
                    ]

                    for pattern in patterns:
                        path_match = re.search(pattern, results_str, re.IGNORECASE)
                        if path_match:
                            csv_path = path_match.group(1).strip().strip('"').strip("'")
                            if os.path.exists(csv_path):
                                st.session_state.clustered_csv_path = csv_path
                                logger.info(f"Found clustered CSV at: {csv_path}")
                                break

                    # If still not found, try to find any CSV file path mentioned
                    if not st.session_state.clustered_csv_path:
                        csv_paths = re.findall(r'["\']?([/\w\-_.]+_clustered\.csv)["\']?', results_str)
                        for path in csv_paths:
                            if os.path.exists(path):
                                st.session_state.clustered_csv_path = path
                                logger.info(f"Found clustered CSV (fallback): {path}")
                                break
            except Exception as e:
                logger.error(f"Could not extract clustered CSV path: {e}")
                st.warning(f"Could not extract clustered CSV path: {e}")

            st.session_state.analysis_complete = True
            # Reset QnA context flag when new analysis is run
            st.session_state.qna_context_initialized = False
            st.session_state.chat_history = []  # Clear chat history for new analysis

            logger.info("Analysis completed successfully")
            st.success("✓ Analysis complete! Download your results below.")

        except Exception as e:
            logger.error(f"Error during analysis: {str(e)}", exc_info=True)
            st.error(f"Error: {e}")
        finally:
            os.remove(tmp_path)
            logger.info("Cleaned up temporary file")

# Step 3: Downloadable results
if st.session_state.clustering_results:
    st.markdown("---")
    st.subheader("📥 Download Results")

    col1, col2, col3 = st.columns(3)

    with col1:
        doc = Document()
        doc.add_heading('Clustering Analysis Results', 0)
        doc.add_paragraph(str(st.session_state.clustering_results))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Clustering Analysis Results",
            data=buffer,
            file_name="clustering_analysis_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with col2:
        if st.session_state.marketing_strategies:
            doc = Document()
            doc.add_heading('Marketing Strategies', 0)
            doc.add_paragraph(str(st.session_state.marketing_strategies))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📊 Marketing Strategies",
                data=buffer,
                file_name="marketing_strategies.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.info("Marketing strategies not captured")

    with col3:
        if st.session_state.clustered_csv_path and os.path.exists(st.session_state.clustered_csv_path):
            with open(st.session_state.clustered_csv_path, 'rb') as f:
                csv_data = f.read()

            st.download_button(
                label="📁 Raw Data with Clusters (CSV)",
                data=csv_data,
                file_name="data_with_clusters.csv",
                mime="text/csv",
                use_container_width=True
            )
        else:
            st.info("CSV with clusters not available")

# Step 4: Q&A Chatbot
if st.session_state.analysis_complete:
    st.markdown("---")
    st.subheader("💬 Q&A Chatbot")

    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    if prompt := st.chat_input("Ask a question about the segmentation results..."):
        logger.info(f"User asked question: {prompt[:50]}...")

        # Add user message to chat history
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                try:
                    # Initialize context on first question only
                    if not st.session_state.qna_context_initialized:
                        logger.info("Initializing QnA context for first question")

                        # Build comprehensive context with the analysis results
                        context_parts = []

                        if st.session_state.clustering_results:
                            context_parts.append(
                                "=== CLUSTERING ANALYSIS RESULTS ===\n" +
                                str(st.session_state.clustering_results)
                            )

                        if st.session_state.marketing_strategies:
                            context_parts.append(
                                "=== MARKETING STRATEGIES ===\n" +
                                str(st.session_state.marketing_strategies)
                            )

                        if not context_parts:
                            logger.error("No analysis context available for QnA")
                            st.error("No analysis context available. Please run the analysis first.")
                        else:
                            context = "\n\n".join(context_parts)
                            # Send context + first question
                            response = run_qna_query(prompt, context)
                            st.session_state.qna_context_initialized = True
                            st.write(response)
                            st.session_state.chat_history.append({"role": "assistant", "content": response})
                            logger.info("QnA context initialized and first response generated")
                    else:
                        # Context already initialized, just send the question
                        response = run_qna_query(prompt, context=None)
                        st.write(response)
                        st.session_state.chat_history.append({"role": "assistant", "content": response})

                except Exception as e:
                    logger.error(f"Error in QnA: {str(e)}", exc_info=True)
                    st.error(f"Error: {str(e)}")
